import io
import math
from collections.abc import Callable
from pathlib import Path
from typing import TypeVar

from app.integrations.base import (
    CapabilityStatus,
    IntegrationError,
    IntegrationInputError,
    IntegrationUnavailableError,
)


T = TypeVar("T")


class DisabledLLMProvider:
    def status(self) -> CapabilityStatus:
        return CapabilityStatus("llm", "disabled", "disabled_no_external_service", False, False, "未配置真实 LLM")

    def generate(self, prompt: str) -> str:
        raise IntegrationUnavailableError("llm", "disabled", "LLM provider is disabled")


class DisabledOCRProvider:
    def status(self) -> CapabilityStatus:
        return CapabilityStatus("ocr", "disabled", "disabled_no_external_service", False, False, "未配置真实 OCR")

    def extract_text(self, content: bytes, content_type: str) -> str:
        raise IntegrationUnavailableError(
            "ocr",
            "disabled",
            "OCR provider is disabled",
            code="CV_OCR_UNAVAILABLE",
        )


class PlainTextDocumentParser:
    def status(self) -> CapabilityStatus:
        return CapabilityStatus("document_parser", "plain_text_local", "local_text_only", True, False, "仅真实解析 UTF-8 纯文本；PDF/Word 不伪造内容")

    def extract_text(self, content: bytes, content_type: str) -> str:
        if content_type != "text/plain":
            raise IntegrationUnavailableError(
                "document_parser",
                "plain_text_local",
                f"No parser configured for {content_type}",
            )
        try:
            return content.decode("utf-8")
        except UnicodeDecodeError as exc:
            raise IntegrationInputError(
                "document_parser",
                "plain_text_local",
                "Text content is not valid UTF-8",
                code="CV_TEXT_EMPTY",
            ) from exc


class PdfTextDocumentParser:
    def status(self) -> CapabilityStatus:
        try:
            version = self._engine_version()
            available = version is not None
        except Exception:
            version = None
            available = False
        return CapabilityStatus(
            "document_parser",
            "pymupdf",
            "pdf_text_local" if available else "unavailable",
            available,
            False,
            f"PyMuPDF {version}" if version else "PyMuPDF unavailable",
            version=str(version) if version else None,
        )

    def _engine_version(self) -> str | None:
        try:
            import fitz

            return str(fitz.VersionBind) if fitz.VersionBind else None
        except Exception:
            return None

    def extract_text(self, content: bytes, content_type: str) -> tuple[str, int]:
        if content_type != "application/pdf":
            raise IntegrationInputError(
                "document_parser",
                "pymupdf",
                f"No PDF parser configured for {content_type}",
                code="CV_PDF_TEXT_EXTRACTION_FAILED",
            )
        try:
            import fitz
        except Exception as exc:
            raise IntegrationUnavailableError(
                "document_parser",
                "pymupdf",
                "PyMuPDF is not installed",
                code="CV_PDF_TEXT_EXTRACTION_FAILED",
            ) from exc
        try:
            document = fitz.open(stream=content, filetype="pdf")
            text = "\n".join(page.get_text("text") for page in document)
            page_count = document.page_count
            document.close()
        except Exception as exc:
            raise IntegrationInputError(
                "document_parser",
                "pymupdf",
                "PDF text extraction failed",
                code="CV_PDF_TEXT_EXTRACTION_FAILED",
            ) from exc
        return text, page_count


class DocxTextDocumentParser:
    def status(self) -> CapabilityStatus:
        try:
            import docx

            version = getattr(docx, "__version__", "unknown")
            available = True
        except Exception:
            version = None
            available = False
        return CapabilityStatus(
            "document_parser",
            "python_docx",
            "docx_text_local" if available else "unavailable",
            available,
            False,
            f"python-docx {version}" if version else "python-docx unavailable",
            version=str(version) if version else None,
        )

    def extract_text(self, content: bytes, content_type: str) -> str:
        if content_type != "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            raise IntegrationInputError(
                "document_parser",
                "python_docx",
                f"No DOCX parser configured for {content_type}",
                code="CV_DOCX_TEXT_EXTRACTION_FAILED",
            )
        try:
            from docx import Document
        except Exception as exc:
            raise IntegrationUnavailableError(
                "document_parser",
                "python_docx",
                "python-docx is not installed",
                code="CV_DOCX_TEXT_EXTRACTION_FAILED",
            ) from exc
        try:
            document = Document(io.BytesIO(content))
            parts = [paragraph.text for paragraph in document.paragraphs]
            for table in document.tables:
                for row in table.rows:
                    parts.extend(cell.text for cell in row.cells)
        except Exception as exc:
            raise IntegrationInputError(
                "document_parser",
                "python_docx",
                "DOCX text extraction failed",
                code="CV_DOCX_TEXT_EXTRACTION_FAILED",
            ) from exc
        return "\n".join(parts)


class TesseractOCRProvider:
    OCR_LANG = "chi_sim+eng"

    def __init__(self) -> None:
        self._version = self._detect_version()

    def _detect_version(self) -> str | None:
        try:
            import pytesseract

            return str(pytesseract.get_tesseract_version())
        except Exception:
            return None

    def status(self) -> CapabilityStatus:
        available = self._version is not None
        return CapabilityStatus(
            "ocr",
            "tesseract",
            "tesseract_cpu_local" if available else "unavailable",
            available,
            False,
            f"Tesseract {self._version}" if available else "Tesseract engine unavailable",
            version=self._version,
        )

    def extract_text(self, content: bytes, content_type: str) -> str:
        text, _ = self.extract_document(content, content_type)
        return text

    def extract_document(
        self, content: bytes, content_type: str
    ) -> tuple[str, list[dict[str, object]]]:
        """Extract ordered text plus word boxes for review highlighting.

        Tesseract automatic page segmentation is kept at PSM 3 so multi-column
        resumes retain layout order. Image preprocessing is deterministic and
        never rewrites OCR text after recognition.
        """
        if self._version is None:
            raise IntegrationUnavailableError(
                "ocr",
                "tesseract",
                "Tesseract engine unavailable",
                code="CV_OCR_UNAVAILABLE",
            )
        try:
            import pytesseract
            from PIL import Image
        except Exception as exc:
            raise IntegrationUnavailableError(
                "ocr",
                "tesseract",
                "Tesseract Python dependencies unavailable",
                code="CV_OCR_UNAVAILABLE",
            ) from exc
        try:
            if content_type in {"image/png", "image/jpeg"}:
                image = Image.open(io.BytesIO(content))
                text, layout = self._ocr_image(image, pytesseract, page=1)
            elif content_type == "application/pdf":
                text, layout = self._ocr_pdf(content, pytesseract)
            else:
                raise IntegrationInputError(
                    "ocr",
                    "tesseract",
                    f"Tesseract does not support {content_type}",
                    code="CV_OCR_FAILED",
                )
        except IntegrationError:
            raise
        except Exception as exc:
            raise IntegrationInputError(
                "ocr",
                "tesseract",
                "OCR failed",
                code="CV_OCR_FAILED",
            ) from exc
        if not text.strip():
            raise IntegrationInputError(
                "ocr",
                "tesseract",
                "OCR produced no text",
                code="CV_OCR_FAILED",
            )
        return text, layout

    @staticmethod
    def _prepare_image(image):
        from PIL import ImageOps

        prepared = ImageOps.exif_transpose(image).convert("L")
        prepared = ImageOps.autocontrast(prepared, cutoff=1)
        if prepared.width < 1800:
            scale = 1800 / max(prepared.width, 1)
            prepared = prepared.resize(
                (1800, max(1, round(prepared.height * scale)))
            )
        return prepared

    def _ocr_image(self, image, pytesseract, *, page: int):
        prepared = self._prepare_image(image)
        data = pytesseract.image_to_data(
            prepared,
            lang=self.OCR_LANG,
            config="--psm 3",
            output_type=pytesseract.Output.DICT,
        )
        words: list[dict[str, object]] = []
        line_values: dict[tuple[int, int, int], list[str]] = {}
        line_order: list[tuple[int, int, int]] = []
        for index, raw_text in enumerate(data.get("text", [])):
            value = str(raw_text).strip()
            if not value:
                continue
            key = (
                int(data["block_num"][index]),
                int(data["par_num"][index]),
                int(data["line_num"][index]),
            )
            if key not in line_values:
                line_values[key] = []
                line_order.append(key)
            line_values[key].append(value)
            words.append(
                {
                    "text": value,
                    "left": int(data["left"][index]),
                    "top": int(data["top"][index]),
                    "width": int(data["width"][index]),
                    "height": int(data["height"][index]),
                    "confidence": float(data["conf"][index]),
                    "block": key[0],
                    "paragraph": key[1],
                    "line": key[2],
                }
            )
        text = "\n".join(" ".join(line_values[key]) for key in line_order)
        return text, [
            {
                "page": page,
                "image_width": prepared.width,
                "image_height": prepared.height,
                "words": words,
            }
        ]

    def _ocr_pdf(self, content: bytes, pytesseract):
        import fitz
        from PIL import Image

        document = fitz.open(stream=content, filetype="pdf")
        chunks: list[str] = []
        layouts: list[dict[str, object]] = []
        for page_number, page in enumerate(document, start=1):
            pixmap = page.get_pixmap(dpi=250)
            image = Image.open(io.BytesIO(pixmap.tobytes("png")))
            text, page_layout = self._ocr_image(
                image, pytesseract, page=page_number
            )
            chunks.append(text)
            layouts.extend(page_layout)
        document.close()
        return "\n".join(chunks), layouts


class DeterministicEmbeddingProvider:
    def __init__(self, dimension: int = 16) -> None:
        self.dimension = dimension

    def status(self) -> CapabilityStatus:
        return CapabilityStatus("embedding", "deterministic_local", "rule_based_deterministic", True, False, f"确定性哈希向量，dimension={self.dimension}，不代表语义模型")

    def embed(self, text: str) -> list[float]:
        if not text.strip():
            raise IntegrationInputError("embedding", "deterministic_local", "Text must not be empty")
        values = [((ord(text[index % len(text)]) % 255) / 127.5) - 1 for index in range(self.dimension)]
        norm = math.sqrt(sum(value * value for value in values)) or 1.0
        return [round(value / norm, 8) for value in values]


class InMemoryVectorStore:
    def __init__(self) -> None:
        self._items: dict[str, tuple[list[float], dict]] = {}

    def status(self) -> CapabilityStatus:
        return CapabilityStatus("vector_store", "memory", "in_memory_test_fallback", True, False, "进程内向量检索，不是 Milvus/Qdrant")

    def upsert(self, object_id: str, vector: list[float], metadata: dict | None = None) -> None:
        if not vector:
            raise IntegrationInputError("vector_store", "memory", "Vector must not be empty")
        self._items[object_id] = (list(vector), dict(metadata or {}))

    def search(self, vector: list[float], top_k: int = 10) -> list[dict]:
        if not vector:
            raise IntegrationInputError("vector_store", "memory", "Vector must not be empty")
        query_norm = math.sqrt(sum(value * value for value in vector)) or 1.0
        results = []
        for object_id, (candidate, metadata) in self._items.items():
            if len(candidate) != len(vector):
                continue
            candidate_norm = math.sqrt(sum(value * value for value in candidate)) or 1.0
            score = sum(a * b for a, b in zip(vector, candidate)) / (query_norm * candidate_norm)
            results.append({"object_id": object_id, "score": round(score, 8), "metadata": metadata})
        return sorted(results, key=lambda item: item["score"], reverse=True)[: max(top_k, 0)]

    def clear(self) -> None:
        self._items.clear()


class DatabaseSyncTaskQueue:
    def status(self) -> CapabilityStatus:
        return CapabilityStatus("task_queue", "database_sync", "database_persisted_sync_executor", True, True, "TaskRecord 持久化；请求内同步执行，不是 Celery/RQ")

    def execute(self, task_type: str, payload: dict, handler: Callable[[dict], T]) -> T:
        return handler(payload)


class LocalFileStorage:
    def __init__(self, root: str | Path) -> None:
        self.root = Path(root)

    def status(self) -> CapabilityStatus:
        return CapabilityStatus("file_storage", "local", "local_filesystem", True, True, "本地单机文件存储，不是对象存储")

    def _path(self, key: str) -> Path:
        if not key or Path(key).name != key:
            raise IntegrationInputError("file_storage", "local", "Storage key must be a plain filename")
        return self.root / key

    def save(self, key: str, content: bytes) -> str:
        path = self._path(key)
        self.root.mkdir(parents=True, exist_ok=True)
        path.write_bytes(content)
        return key

    def read(self, key: str) -> bytes:
        path = self._path(key)
        if not path.exists():
            raise IntegrationInputError("file_storage", "local", "Storage object not found")
        return path.read_bytes()

    def delete(self, key: str) -> None:
        self._path(key).unlink(missing_ok=True)


class KeywordEvidenceRetriever:
    def status(self) -> CapabilityStatus:
        return CapabilityStatus("evidence_retriever", "keyword_local", "rule_based_keyword_retrieval", True, False, "关键词重合降级检索，不是向量检索")

    def retrieve(self, query: str, documents: list[dict], top_k: int = 5) -> list[dict]:
        terms = {term.lower() for term in query.split() if term.strip()}
        scored = []
        for document in documents:
            text = str(document.get("text") or document.get("raw_text") or "").lower()
            score = sum(1 for term in terms if term in text)
            if score:
                scored.append({**document, "rule_score": score})
        return sorted(scored, key=lambda item: item["rule_score"], reverse=True)[: max(top_k, 0)]


class DisabledTrendSourceCrawler:
    def status(self) -> CapabilityStatus:
        return CapabilityStatus("trend_crawler", "disabled", "disabled_no_network", False, False, "默认测试和开发不访问真实网络")

    def fetch(self, source: dict) -> list[dict]:
        raise IntegrationUnavailableError("trend_crawler", "disabled", "Trend source crawler is disabled")
